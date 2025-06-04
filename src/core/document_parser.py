import time
import re
from pathlib import Path
from typing import Optional, List, Dict, Any, Union
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.shared import Inches

from src.models.parsing import ParsedDocument, DocumentPage, TextBlock
from src.utils.logging import get_logger
from src.utils.file_utils import get_file_info

logger = get_logger(__name__)


class DocumentParsingError(Exception):
    """Custom exception for document parsing errors."""
    pass


class DocumentParser:
    """
    Parse PDF and Word documents using PyMuPDF and python-docx.
    """
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.doc'}
    
    def parse_document(self, file_path: str) -> ParsedDocument:
        """
        Parse a document and return structured content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ParsedDocument with content and metadata
        """
        start_time = time.time()
        path = Path(file_path)
        
        if not path.exists():
            raise DocumentParsingError(f"File not found: {file_path}")
        
        file_type = path.suffix.lower().lstrip('.')
        
        if file_type not in {'pdf', 'docx', 'doc'}:
            raise DocumentParsingError(f"Unsupported file type: {file_type}")
        
        try:
            if file_type == 'pdf':
                parsed_doc = self._parse_pdf(file_path)
            elif file_type in ['docx', 'doc']:
                parsed_doc = self._parse_docx(file_path)
            else:
                raise DocumentParsingError(f"Unsupported file type: {file_type}")
            
            # Calculate parsing duration
            parsing_duration_ms = int((time.time() - start_time) * 1000)
            parsed_doc.parsing_duration_ms = parsing_duration_ms
            
            logger.info(f"Successfully parsed {file_path} in {parsing_duration_ms}ms")
            return parsed_doc
            
        except Exception as e:
            parsing_duration_ms = int((time.time() - start_time) * 1000)
            logger.error(f"Failed to parse {file_path}: {e}", exc_info=True)
            
            return ParsedDocument(
                file_path=file_path,
                file_type=file_type,
                parsing_method='pymupdf' if file_type == 'pdf' else 'python_docx',
                parsing_success=False,
                parsing_error_message=str(e),
                parsing_duration_ms=parsing_duration_ms
            )
    
    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """
        Parse PDF document using PyMuPDF.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            ParsedDocument with extracted content
        """
        doc = fitz.open(file_path)
        
        try:
            # Extract document metadata
            doc_metadata = doc.metadata
            
            # Initialize content containers
            all_text = []
            all_markdown = []
            page_count = len(doc)
            
            # Process each page
            for page_num in range(page_count):
                page = doc.load_page(page_num)
                page_content = self._extract_page_content(page, page_num + 1)
                
                all_text.append(page_content['text'])
                all_markdown.append(page_content['markdown'])
            
            # Combine all content
            raw_text = '\n\n'.join(all_text)
            markdown_content = '\n\n'.join(all_markdown)
            
            # Calculate word count
            word_count = len(raw_text.split()) if raw_text else 0
            
            return ParsedDocument(
                file_path=file_path,
                file_type='pdf',
                parsing_method='pymupdf',
                parsing_success=True,
                markdown_content=markdown_content,
                raw_text=raw_text,
                page_count=page_count,
                word_count=word_count,
                document_metadata=doc_metadata
            )
            
        finally:
            doc.close()
    
    def _extract_page_content(self, page: fitz.Page, page_num: int) -> Dict[str, str]:
        """
        Extract content from a single PDF page.
        
        Args:
            page: PyMuPDF page object
            page_num: Page number (1-indexed)
            
        Returns:
            Dictionary with text and markdown content
        """
        # Get text blocks with formatting information
        blocks = page.get_text("dict")
        
        text_content = []
        markdown_content = []
        
        for block in blocks["blocks"]:
            if "lines" not in block:
                continue
            
            for line in block["lines"]:
                line_text = ""
                line_markdown = ""
                
                for span in line["spans"]:
                    text = span["text"].strip()
                    if not text:
                        continue
                    
                    # Basic formatting detection
                    font_size = span.get("size", 12)
                    flags = span.get("flags", 0)
                    
                    # Check for bold/italic
                    is_bold = bool(flags & 2**4)
                    is_italic = bool(flags & 2**1)
                    
                    # Detect headings based on font size
                    if font_size > 14:
                        # Likely a heading
                        heading_level = min(3, max(1, int((font_size - 12) / 2)))
                        line_markdown += "#" * heading_level + " " + text + "\n\n"
                    else:
                        # Regular text with formatting
                        formatted_text = text
                        if is_bold and is_italic:
                            formatted_text = f"***{text}***"
                        elif is_bold:
                            formatted_text = f"**{text}**"
                        elif is_italic:
                            formatted_text = f"*{text}*"
                        
                        line_markdown += formatted_text + " "
                    
                    line_text += text + " "
                
                if line_text.strip():
                    text_content.append(line_text.strip())
                    if not line_markdown.endswith("\n\n"):
                        markdown_content.append(line_markdown.strip())
        
        # Handle tables
        tables = page.find_tables()
        for table in tables:
            table_data = table.extract()
            if table_data:
                # Convert table to markdown
                markdown_table = self._format_table_as_markdown(table_data)
                markdown_content.append(markdown_table)
        
        # Join content
        page_text = '\n'.join(text_content)
        page_markdown = '\n'.join(markdown_content)
        
        # Add page break
        if page_markdown:
            page_markdown += f"\n\n---\n*Page {page_num}*\n\n"
        
        return {
            'text': page_text,
            'markdown': page_markdown
        }
    
    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """
        Parse Word document using python-docx.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            ParsedDocument with extracted content
        """
        try:
            doc = DocxDocument(file_path)
            
            # Extract document properties
            doc_metadata = {
                'title': doc.core_properties.title,
                'author': doc.core_properties.author,
                'subject': doc.core_properties.subject,
                'created': doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                'modified': doc.core_properties.modified.isoformat() if doc.core_properties.modified else None
            }
            
            # Extract content
            text_content = []
            markdown_content = []
            
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    continue
                
                # Detect heading style
                if paragraph.style.name.startswith('Heading'):
                    try:
                        heading_level = int(paragraph.style.name.split()[-1])
                        heading_level = min(6, max(1, heading_level))
                    except (ValueError, IndexError):
                        heading_level = 1
                    
                    markdown_content.append(f"{'#' * heading_level} {paragraph.text}\n")
                else:
                    # Regular paragraph
                    para_text = paragraph.text
                    para_markdown = self._format_paragraph_markdown(paragraph)
                    
                    text_content.append(para_text)
                    markdown_content.append(para_markdown)
            
            # Handle tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    markdown_table = self._format_table_as_markdown(table_data)
                    markdown_content.append(markdown_table)
            
            # Combine content
            raw_text = '\n'.join(text_content)
            markdown_text = '\n'.join(markdown_content)
            
            # Calculate statistics
            word_count = len(raw_text.split()) if raw_text else 0
            
            return ParsedDocument(
                file_path=file_path,
                file_type='docx',
                parsing_method='python_docx',
                parsing_success=True,
                markdown_content=markdown_text,
                raw_text=raw_text,
                page_count=1,  # Word docs don't have explicit pages in python-docx
                word_count=word_count,
                document_metadata=doc_metadata
            )
            
        except Exception as e:
            raise DocumentParsingError(f"Failed to parse DOCX file: {e}")
    
    def _format_paragraph_markdown(self, paragraph) -> str:
        """
        Format a Word paragraph as markdown with inline formatting.
        
        Args:
            paragraph: python-docx paragraph object
            
        Returns:
            Markdown formatted text
        """
        markdown_text = ""
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Apply formatting
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            
            markdown_text += text
        
        return markdown_text + "\n"
    
    def _format_table_as_markdown(self, table_data: List[List[str]]) -> str:
        """
        Format table data as markdown table.
        
        Args:
            table_data: List of rows, each row is a list of cell values
            
        Returns:
            Markdown formatted table
        """
        if not table_data:
            return ""
        
        # Find max width for each column
        col_count = max(len(row) for row in table_data) if table_data else 0
        col_widths = [0] * col_count
        
        # Normalize rows to same length
        normalized_rows = []
        for row in table_data:
            normalized_row = row + [''] * (col_count - len(row))
            normalized_rows.append(normalized_row)
            
            # Update column widths
            for i, cell in enumerate(normalized_row):
                col_widths[i] = max(col_widths[i], len(str(cell)))
        
        # Build markdown table
        markdown_lines = []
        
        for i, row in enumerate(normalized_rows):
            # Format row
            formatted_cells = []
            for j, cell in enumerate(row):
                formatted_cells.append(str(cell).ljust(col_widths[j]))
            
            markdown_lines.append("| " + " | ".join(formatted_cells) + " |")
            
            # Add separator after header row
            if i == 0:
                separator_cells = ['-' * width for width in col_widths]
                markdown_lines.append("|-" + "-|-".join(separator_cells) + "-|")
        
        return '\n'.join(markdown_lines) + '\n\n'
    
    def validate_parsing_quality(self, parsed_doc: ParsedDocument) -> Dict[str, Any]:
        """
        Validate the quality of parsed content.
        
        Args:
            parsed_doc: Parsed document to validate
            
        Returns:
            Dictionary with quality metrics
        """
        if not parsed_doc.parsing_success:
            return {
                'is_valid': False,
                'error': parsed_doc.parsing_error_message,
                'word_count': 0,
                'content_ratio': 0.0
            }
        
        raw_text = parsed_doc.raw_text or ""
        markdown_content = parsed_doc.markdown_content or ""
        
        # Basic quality checks
        word_count = len(raw_text.split())
        char_count = len(raw_text)
        
        # Check for minimum content
        has_minimum_content = word_count >= 10 and char_count >= 50
        
        # Check content/markup ratio
        content_ratio = len(raw_text) / max(len(markdown_content), 1)
        
        # Check for excessive repetition
        words = raw_text.split()
        unique_words = set(words)
        repetition_ratio = len(unique_words) / max(len(words), 1)
        
        is_valid = (
            has_minimum_content and
            content_ratio > 0.3 and  # Not too much markup
            repetition_ratio > 0.1    # Not too repetitive
        )
        
        return {
            'is_valid': is_valid,
            'word_count': word_count,
            'char_count': char_count,
            'content_ratio': content_ratio,
            'repetition_ratio': repetition_ratio,
            'has_minimum_content': has_minimum_content
        }